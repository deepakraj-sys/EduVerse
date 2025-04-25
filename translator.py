import streamlit as st
from deep_translator import GoogleTranslator
import PyPDF2  # Replaced fitz (PyMuPDF)
import docx
import io
import pyttsx3
import speech_recognition as sr
from transformers import pipeline
from docx import Document

# Set page config
st.set_page_config(page_title="Real-Time Translator", layout="centered")
st.title("🌍 Real-time Multilingual Translator")

# Preferred languages
preferred_langs = [
    "english", "tamil", "hindi", "telugu", "kannada", "malayalam",
    "bengali", "oriya", "arabic", "french"
]

# Load supported languages
langs = GoogleTranslator().get_supported_languages(as_dict=True)
lang_keys = list(langs.keys())
available_preferred = [lang for lang in preferred_langs if lang in lang_keys]
final_lang_list = available_preferred + [lang for lang in lang_keys if lang not in available_preferred]

# Language selection
col1, col2 = st.columns(2)
with col1:
    source_lang = st.selectbox("Source Language (select 'auto' to auto-detect)", ["auto"] + final_lang_list, index=0)
with col2:
    target_lang = st.selectbox("Target Language", final_lang_list, index=final_lang_list.index("tamil"))

# File reader
def read_file(file):
    try:
        if file.name.endswith(".txt"):
            return str(file.read(), "utf-8")
        elif file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        elif file.name.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file.read()))
            return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"File reading failed: {e}")
    return ""

# Summarizer model
@st.cache_resource
def load_summarizer():
    return pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

summarizer = load_summarizer()

# State variables
if "input_text" not in st.session_state:
    st.session_state.input_text = ""

# Mic input
if st.button("🎙️ Use Microphone"):
    r = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("🎧 Listening...")
        try:
            audio = r.listen(source, timeout=5)
            st.session_state.input_text = sr.Recognize_google(audio)
            st.success(f"You said: {st.session_state.input_text}")
        except Exception as e:
            st.error(f"Microphone error: {e}")

# File upload
uploaded_file = st.file_uploader("📁 Upload TXT, PDF, or DOCX", type=["txt", "pdf", "docx"])
if uploaded_file:
    file_text = read_file(uploaded_file)
    st.session_state.input_text = file_text
    st.success("✅ File content loaded.")

# Text input area
st.session_state.input_text = st.text_area("✍️ Enter your prompt or content:", value=st.session_state.input_text, height=200)

# Translate
translated_text = ""
if st.button("🔄 Translate"):
    if st.session_state.input_text.strip() == "":
        st.warning("⚠️ Please enter some text first.")
    else:
        try:
            translated_text = GoogleTranslator(source=source_lang, target=target_lang).translate(st.session_state.input_text)
            st.success("✅ Translation complete!")
            st.text_area("📝 Translated Output", value=translated_text, height=300)
        except Exception as e:
            st.error(f"❌ Translation error: {e}")

# Speak Output
if translated_text and st.button("🔊 Speak Output"):
    try:
        tts_engine = pyttsx3.init()
        tts_engine.say(translated_text)
        tts_engine.runAndWait()
    except Exception as e:
        st.error(f"TTS error: {e}")

# Summarize + Translate
if st.button("🧠 Summarize + Translate"):
    try:
        summary = summarizer(st.session_state.input_text, max_length=100, min_length=30, do_sample=False)[0]['summary_text']
        translated_summary = GoogleTranslator(source="auto", target=target_lang).translate(summary)
        st.success("🧾 Summary & Translation complete!")
        st.text_area("🧠 Summary + Translation", value=translated_summary, height=200)
    except Exception as e:
        st.error(f"Summary/translation error: {e}")

# Download options
if translated_text:
    col1, col2 = st.columns(2)

    with col1:
        st.download_button("📥 Download as PDF", data=translated_text.encode('utf-8'),
                           file_name="translated_output.pdf", mime="application/pdf")

    with col2:
        def generate_docx(text):
            doc_io = io.BytesIO()
            document = Document()
            document.add_paragraph(text)
            document.save(doc_io)
            doc_io.seek(0)
            return doc_io

        docx_io = generate_docx(translated_text)
        st.download_button("📥 Download as Word (DOCX)", data=docx_io,
                           file_name="translated_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")